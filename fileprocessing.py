import os, tempfile #, ocrmypdf, openpyxl, fitz, pytesseract, io
import pandas as pd
import streamlit as st
from streamlit.runtime.uploaded_file_manager import UploadedFile
from docx import Document as DocxDocument
from PIL import Image
# Langchain imports
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Process documents for the user: import os, tempfile, pandas, langchain_text_splitters, langchain_core.documents, PyMuPDFLoader, Document
def process_document(uploaded_file: UploadedFile) -> tuple:
    """
    Process various document types and split them into chunks for vector storage.
    
    Args:
        uploaded_file: The uploaded file to process
        
    Returns:
        A tuple containing (all_splits, file_size_bytes)
    """
    # Store uploaded file as a temp file before it goes to vector database
    file_extension = uploaded_file.name.split(".")[-1].lower()
    temp_file = tempfile.NamedTemporaryFile(suffix=f".{file_extension}", delete=False)
    temp_file.write(uploaded_file.read())
    temp_file.seek(0)
    file_size_bytes = os.path.getsize(temp_file.name)
    temp_file.close()
    documents = []

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500, chunk_overlap=100, separators=["\n\n", "\n", ".", "?", "!", " ", ""]
    )

    try:
        if file_extension == "docx":
            doc = DocxDocument(temp_file.name)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            documents.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
        elif file_extension == "pdf":
            try:
                # First try using OCRMyPDF to convert scanned images to searchable PDF
                loader = PyMuPDFLoader(temp_file.name)
                documents = loader.load()
                if not any(doc.page_content.strip() for doc in documents):
                    raise Exception("OCRMyPDF produced empty text")
            except Exception as e:
                # If OCRMyPDF fails, try manual OCR approach
                st.error(f"Failed to read document, maybe this is a scanned pdf file. {str(e)}")
        elif file_extension == "txt":
            with open(temp_file.name, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            documents.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
        elif file_extension in ["xls", "xlsm", "xlsx"]: 
            excel_df = pd.read_excel(temp_file.name, sheet_name=None)
            for sheet_name, sheet_df in excel_df.items():
                text = sheet_df.to_csv(index=False)
                documents.append(Document(page_content=text, metadata={"source": f"{uploaded_file.name} - {sheet_name}"}))
        elif file_extension == "csv":
            csv_df = pd.read_csv(temp_file.name)
            text = csv_df.to_csv(index=False)
            documents.append(Document(page_content=text, metadata = {"source": uploaded_file.name}))
        
        elif file_extension in ["py", "r", "js", "jsx", "ts", "tsx", "html", "htm", "xml", "css", "scss", "less", 
                                "java", "c", "h", "cpp", "hpp", "cxx", "cc", "cs", "php", "rb", "go", "sh", 
                                "bash", "swift", "kt", "kts", "rs", "scala", "pl", "pm", "sql", "yaml", "yml"]:
            # Read code files as plain text
            with open(temp_file.name, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            documents.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
        else:
            st.error(f"Unsupported file format: {file_extension}. Please upload a PDF, DOCX, TXT, Excel, or CSV file.")
            os.unlink(temp_file.name)
            return [], 0
    except Exception as e:
        st.error(f"Error processing {file_extension} file: {str(e)}")
        os.unlink(temp_file.name)
        return [], 0

    # Split documents into chunks
    all_splits = []
    if documents:
        splits_raw = text_splitter.split_documents(documents)
        # Extract content and filter out empty chunks
        all_splits = [Document(page_content=split.page_content.strip(), metadata=split.metadata) 
                    for split in splits_raw if split.page_content.strip()]
    
    # Clean up the temp file
    os.unlink(temp_file.name)
    
    return all_splits, file_size_bytes

    